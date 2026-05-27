"""ER Diagram views for YanLeafAdmin"""
import json
from django.contrib.admin.views.decorators import staff_member_required
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render
from django.apps import apps as django_apps
from django.db.models.fields import Field
from django.db.models.fields.related import ForeignKey, OneToOneField, RelatedField
from .sql_parser import SQLParser


@staff_member_required
def er_diagram_page(request):
    """ER diagram main page"""
    return render(request, 'erd/er_diagram.html', {
        'title': '数据库 ER 图',
    })


@staff_member_required
def er_model_api(request):
    """Return ER data from Django models as JSON"""
    nodes = []
    edges = []
    table_ids = {}

    for model in django_apps.get_models():
        app_label = model._meta.app_label
        model_name = model._meta.model_name
        table_name = model._meta.db_table
        table_node_id = f"tbl_{app_label}_{model_name}"

        if table_node_id in table_ids:
            continue
        table_ids[table_name] = table_node_id

        # Collect columns
        columns = []
        for field in model._meta.get_fields():
            # Skip non-field relation objects (e.g. ManyToManyRel, ManyToOneRel)
            if not isinstance(field, Field):
                continue
            col_info = {
                'name': field.name,
                'type': field.get_internal_type(),
                'pk': getattr(field, 'primary_key', False),
                'fk': isinstance(field, (ForeignKey, OneToOneField)),
            }
            if field.primary_key:
                col_info['pk'] = True
            if isinstance(field, (ForeignKey, OneToOneField)):
                col_info['ref_table'] = field.related_model._meta.db_table
                col_info['ref_column'] = field.target_field.name
            # Get verbose name
            if hasattr(field, 'verbose_name'):
                col_info['verbose'] = str(field.verbose_name)
            columns.append(col_info)

        nodes.append({
            'id': table_node_id,
            'name': model._meta.verbose_name or model_name,
            'table_name': table_name,
            'columns': columns,
            'app': app_label,
        })

    # Build edges from FK relationships
    edge_ids = set()
    for model in django_apps.get_models():
        app_label = model._meta.app_label
        model_name = model._meta.model_name
        src_id = f"tbl_{app_label}_{model_name}"

        for field in model._meta.get_fields():
            if not isinstance(field, (ForeignKey, OneToOneField)):
                continue
            ref_table = field.related_model._meta.db_table
            if ref_table in table_ids:
                dst_id = table_ids[ref_table]
                edge_key = f"{src_id}->{dst_id}"
                if edge_key not in edge_ids:
                    edge_ids.add(edge_key)
                    edges.append({
                        'source': src_id,
                        'target': dst_id,
                        'label': field.name,
                    })

    return JsonResponse({'nodes': nodes, 'edges': edges})


@staff_member_required
def er_sql_api(request):
    """Parse user-submitted SQL DDL and return ER data as JSON"""
    sql_text = request.POST.get('sql', '') or request.GET.get('sql', '')
    if not sql_text:
        return JsonResponse({'error': '请提供 SQL DDL 语句'}, status=400)

    try:
        parser = SQLParser()
        tables = parser.parse(sql_text)
    except Exception as e:
        return JsonResponse({'error': f'SQL 解析失败: {str(e)}'}, status=400)

    nodes = []
    edges = []
    edge_ids = set()
    table_names = {t.name.lower() for t in tables}

    for table in tables:
        node_id = f"tbl_{table.name}"
        columns = [c.to_dict() for c in table.columns]
        nodes.append({
            'id': node_id,
            'name': table.comment or table.name,
            'table_name': table.name,
            'columns': columns,
        })

        for col in table.columns:
            if col.is_foreign and col.reference_table:
                ref_lower = col.reference_table.lower()
                if ref_lower in table_names:
                    dst_id = f"tbl_{ref_lower}"
                    edge_key = f"{node_id}->{dst_id}"
                    if edge_key not in edge_ids:
                        edge_ids.add(edge_key)
                        edges.append({
                            'source': node_id,
                            'target': dst_id,
                            'label': col.name,
                        })

    return JsonResponse({'nodes': nodes, 'edges': edges})


@staff_member_required
def er_word_export(request):
    """Generate Word .docx with three-line tables"""
    tables_json = request.POST.get('tables', '')
    if not tables_json:
        return HttpResponse('No data', status=400)
    try:
        table_list = json.loads(tables_json)
    except (json.JSONDecodeError, TypeError):
        return HttpResponse('Invalid data', status=400)

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return HttpResponse('Word 导出需要安装 python-docx：pip install yanleafadmin[word]', status=500)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    for t in table_list:
        title = (t.get('comment') or t.get('name', ''))
        h = doc.add_heading(title + ' (' + t.get('name', '') + ')', level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cols = t.get('columns', [])
        table = doc.add_table(rows=len(cols) + 1, cols=4)
        table.style = 'Table Grid'

        for j, hd in enumerate(['字段名', '类型', '约束', '说明']):
            cell = table.rows[0].cells[j]
            cell.text = hd
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for i, col in enumerate(cols):
            row = table.rows[i + 1]
            row.cells[0].text = col.get('name', '')
            row.cells[1].text = col.get('data_type', col.get('type', ''))
            cons = []
            if col.get('is_primary') or col.get('pk'):
                cons.append('PK')
            if col.get('is_foreign') or col.get('fk'):
                ref = col.get('reference_table', col.get('ref_table', ''))
                cons.append('FK->' + ref if ref else 'FK')
            row.cells[2].text = ', '.join(cons)
            row.cells[3].text = col.get('comment', col.get('verbose', ''))
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        tbl = table._tbl
        borders = OxmlElement('w:tblBorders')
        tbl.insert(0, borders)
        for bn, sz in [('top', 12), ('bottom', 12), ('insideH', 4)]:
            el = OxmlElement('w:' + bn)
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)

        doc.add_paragraph('')

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="er-tables.docx"'
    doc.save(response)
    return response
